"""
cv/cv_parser.py — Parses uploaded master CV into named sections.

Supports .tex, .pdf, and .docx input formats (auto-detected by extension).
Returns a dict mapping section names to section text.
"""

import logging
import re
from pathlib import Path

logger = logging.getLogger(__name__)

# Common resume section names used for boundary detection in PDF/DOCX
_COMMON_SECTIONS = {
    "education", "experience", "work experience", "employment",
    "projects", "skills", "technical skills", "publications",
    "certifications", "awards", "summary", "objective",
    "professional experience", "research", "interests",
    "achievements", "activities", "volunteer", "references",
    "header", "contact", "about", "profile",
}


def parse_cv(cv_path: Path) -> dict[str, str]:
    """Parse a CV file into named sections. Returns {section_name: section_text}."""
    cv_path = Path(cv_path)

    if not cv_path.exists():
        raise FileNotFoundError(f"CV file not found: {cv_path}")

    suffix = cv_path.suffix.lower()

    if suffix == ".tex":
        return _parse_tex(cv_path)
    elif suffix == ".pdf":
        return _parse_pdf(cv_path)
    elif suffix == ".docx":
        return _parse_docx(cv_path)
    else:
        raise ValueError(
            f"Unsupported CV format: '{suffix}'. "
            f"Supported formats: .tex, .pdf, .docx"
        )


def _parse_tex(cv_path: Path) -> dict[str, str]:
    """Parse a LaTeX CV by section/subsection commands or explicit markers."""
    text = cv_path.read_text(encoding="utf-8")
    sections: dict[str, str] = {}

    # Pattern 1: Explicit markers  %--- SECTION: <NAME> ---
    marker_pattern = re.compile(r"%---\s*SECTION:\s*(.+?)\s*---")
    # Pattern 2: LaTeX section commands  \section{Name} or \subsection{Name}
    latex_pattern = re.compile(r"\\(?:sub)?section\*?\{([^}]+)\}")

    # Combine both patterns — find all section boundaries
    boundaries: list[tuple[int, str]] = []

    for match in marker_pattern.finditer(text):
        boundaries.append((match.start(), match.group(1).strip()))

    for match in latex_pattern.finditer(text):
        boundaries.append((match.end(), match.group(1).strip()))

    if not boundaries:
        logger.warning("No section boundaries found in %s — returning full text", cv_path)
        return {"full_document": text.strip()}

    # Sort by position in file
    boundaries.sort(key=lambda x: x[0])

    for i, (pos, name) in enumerate(boundaries):
        if i + 1 < len(boundaries):
            end_pos = boundaries[i + 1][0]
            # For the next boundary, find the start of that line to avoid
            # including the next section command in this section's text
            line_start = text.rfind("\n", 0, end_pos)
            if line_start == -1:
                line_start = end_pos
            content = text[pos:line_start]
        else:
            content = text[pos:]

        # Clean up: remove the section command itself from content start
        # (for latex_pattern matches, pos is already after the command)
        content = content.strip()

        # Remove trailing \end{document} if present
        content = re.sub(r"\\end\{document\}\s*$", "", content).strip()

        if content:
            sections[name] = content

    logger.info("Parsed %d sections from LaTeX CV: %s", len(sections), list(sections.keys()))
    return sections


def _parse_pdf(cv_path: Path) -> dict[str, str]:
    """Parse a PDF CV by extracting text and splitting on section-like headings."""
    try:
        import pdfplumber
    except ImportError as e:
        raise ImportError(
            "pdfplumber is required for PDF parsing. "
            "Install it: pip install pdfplumber"
        ) from e

    pages_text: list[str] = []
    with pdfplumber.open(cv_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                pages_text.append(page_text)

    if not pages_text:
        logger.warning("No text extracted from PDF: %s", cv_path)
        return {}

    full_text = "\n".join(pages_text)
    return _split_by_headings(full_text, cv_path)


def _parse_docx(cv_path: Path) -> dict[str, str]:
    """Parse a DOCX CV by Heading styles or ALL-CAPS line detection."""
    try:
        import docx
    except ImportError as e:
        raise ImportError(
            "python-docx is required for DOCX parsing. "
            "Install it: pip install python-docx"
        ) from e

    doc = docx.Document(cv_path)

    # First try: use Heading styles as section boundaries
    sections: dict[str, str] = {}
    current_section: str | None = None
    current_lines: list[str] = []

    has_headings = any(
        p.style.name.startswith("Heading") for p in doc.paragraphs
    )

    if has_headings:
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading") and para.text.strip():
                # Save previous section
                if current_section is not None:
                    sections[current_section] = "\n".join(current_lines).strip()
                current_section = para.text.strip()
                current_lines = []
            else:
                if para.text.strip():
                    current_lines.append(para.text.strip())

        if current_section is not None:
            sections[current_section] = "\n".join(current_lines).strip()

        # Remove empty sections
        sections = {k: v for k, v in sections.items() if v}

        if sections:
            logger.info(
                "Parsed %d sections from DOCX (Heading styles): %s",
                len(sections), list(sections.keys()),
            )
            return sections

    # Fallback: extract all text and split by headings
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    return _split_by_headings(full_text, cv_path)


def _split_by_headings(text: str, source_path: Path) -> dict[str, str]:
    """Split plain text into sections using ALL-CAPS lines or known section names."""
    lines = text.split("\n")
    sections: dict[str, str] = {}
    current_section: str | None = None
    current_lines: list[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_lines.append("")
            continue

        if _is_section_heading(stripped):
            # Save previous section
            if current_section is not None:
                sections[current_section] = "\n".join(current_lines).strip()
            current_section = stripped.title()
            current_lines = []
        else:
            current_lines.append(stripped)

    # Save last section
    if current_section is not None:
        sections[current_section] = "\n".join(current_lines).strip()

    # Remove empty sections
    sections = {k: v for k, v in sections.items() if v}

    if not sections:
        logger.warning(
            "No section headings detected in %s — returning full text",
            source_path,
        )
        return {"full_document": text.strip()}

    logger.info(
        "Parsed %d sections from text: %s", len(sections), list(sections.keys()),
    )
    return sections


def _is_section_heading(line: str) -> bool:
    """Detect if a line is likely a section heading."""
    # ALL-CAPS line with 2+ alpha chars and no more than 5 words
    words = line.split()
    if (
        len(words) <= 5
        and len(line) >= 2
        and line.upper() == line
        and any(c.isalpha() for c in line)
        and not line.startswith("•")
        and not line.startswith("-")
    ):
        return True

    # Known section name (case-insensitive)
    normalized = line.lower().rstrip(":").strip()
    if normalized in _COMMON_SECTIONS:
        return True

    return False
